from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

doc.add_heading('Приглашение', 0)
doc.add_heading('Уважаемый(ая) {{ФИО}},', level=1)

doc.add_paragraph('Мы рады пригласить вас на нашу свадьбу.')
doc.add_paragraph('Дата: {{Дата}}')
doc.add_paragraph('Время: {{Время}}')
doc.add_paragraph('Место: {{Место}}')
doc.add_paragraph('Адрес: {{Адрес}}')


doc.add_paragraph('Список гостей', style='Heading2')
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ФИО'
hdr_cells[1].text = 'Время'
hdr_cells[2].text = 'Место'

doc.add_paragraph('План мероприятия:', style='Heading2')
doc.add_paragraph('Регистрация гостей', style='BodyText').runs[0].italic = True
doc.add_paragraph('Церемония', style='BodyText').runs[0].italic = True
doc.add_paragraph('Банкет', style='BodyText').runs[0].italic = True

doc.add_paragraph('Необходимо взять с собой:', style='Heading2')
doc.add_paragraph('Паспорт', style='List Number')
doc.add_paragraph('Подарок', style='List Number')
doc.add_paragraph('Хорошее настроение', style='List Number')

doc.add_picture('signature.png', width=Inches(1.25))

doc.save('invitation.docx')